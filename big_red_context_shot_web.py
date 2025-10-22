#!/usr/bin/env python3
"""
Big Red Button — Cross-Platform Web Edition (Windows-friendly)
--------------------------------------------------------------
- Windows/Linux: uses MSS to capture the full virtual screen.
- macOS: uses `screencapture` (native, no extra permissions beyond Screen Recording).

Workflow:
  click -> wait <delay> seconds -> screenshot -> append to Word table

The Word file: ~/Documents/ContextShots.docx
Two columns per row: [Context + timestamp] | [Screenshot]
Rotates after 90 rows (ContextShots (2).docx, etc.).
"""
import os
import time
import shlex
import platform
import subprocess
from datetime import datetime
from pathlib import Path
import webbrowser

from flask import Flask, request, render_template_string, redirect, url_for, flash

# Word
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Screenshot (Windows/Linux)
from mss import mss
from mss.tools import to_png

APP_TITLE = "Big Red Button - Context + Screenshot (Cross-Platform)"
HOST = "127.0.0.1"
PORT = 8788
MAX_ROWS_PER_DOC = 90
DEFAULT_DOC_NAME = "ContextShots.docx"
DEFAULT_DELAY_SECONDS = 2.0

app = Flask(__name__)
app.secret_key = "context-shot-cross"

def user_documents_dir() -> Path:
    home = Path.home()
    docs = home / "Documents"
    return docs if docs.exists() else home

def next_available_filename(base: Path) -> Path:
    if not base.exists():
        return base
    stem, suffix = base.stem, base.suffix
    n = 2
    while True:
        cand = base.with_name(f"{stem} ({n}){suffix}")
        if not cand.exists():
            return cand
        n += 1

def _set_table_column_widths(table, widths_in_inches):
    for idx, width in enumerate(widths_in_inches):
        for cell in table.columns[idx].cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcW = OxmlElement('w:tcW')
            tcW.set(qn('w:type'), 'dxa')
            tcW.set(qn('w:w'), str(int(width * 1440)))
            tcPr.append(tcW)

def _add_title_and_table(doc: Document, title_suffix: str = ""):
    title = doc.add_paragraph()
    run = title.add_run(f"Context + Screenshot Log {title_suffix}".strip())
    run.font.size = Pt(16)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    table = doc.add_table(rows=1, cols=2)
    hdr = table.rows[0].cells
    hdr[0].text = "Context (with timestamp)"
    hdr[1].text = "Screenshot"
    _set_table_column_widths(table, [3.1, 3.1])

def ensure_document_and_table(doc_path: Path):
    if doc_path.exists():
        doc = Document(str(doc_path))
        if doc.tables:
            table = doc.tables[0]
            if len(table.columns) != 2:
                doc = Document()
                _add_title_and_table(doc)
                return doc, doc.tables[0], doc_path
            rows = max(0, len(table.rows) - 1)
            if rows >= MAX_ROWS_PER_DOC:
                new_path = next_available_filename(doc_path)
                doc = Document()
                _add_title_and_table(doc, title_suffix=f"(Part {new_path.stem.split()[-1].strip('()')})")
                return doc, doc.tables[0], new_path
            return doc, table, doc_path
        else:
            _add_title_and_table(doc)
            return doc, doc.tables[0], doc_path
    else:
        doc = Document()
        _add_title_and_table(doc)
        return doc, doc.tables[0], doc_path

def compute_column_image_width_inches(doc: Document) -> float:
    section = doc.sections[0]
    page_width = section.page_width / 914400
    left = section.left_margin / 914400
    right = section.right_margin / 914400
    usable = page_width - left - right
    return max(2.2, min((usable / 2.0) - 0.15, 3.5))

def take_full_screenshot_to(file_path: Path) -> None:
    system = platform.system().lower()
    if system == "darwin":
        cmd = f'screencapture -x -t png {shlex.quote(str(file_path))}'
        subprocess.run(cmd, shell=True, check=True)
        return
    with mss() as sct:
        monitor = sct.monitors[0]
        raw = sct.grab(monitor)
        to_png(raw.rgb, raw.size, output=str(file_path))

def append_entry(doc_path: Path, context_text: str, screenshot_path: Path) -> Path:
    target = doc_path
    if target.exists():
        doc_tmp = Document(str(target))
        table_tmp = doc_tmp.tables[0] if doc_tmp.tables else None
        rows_tmp = (len(table_tmp.rows) - 1) if table_tmp else 0
        if rows_tmp >= MAX_ROWS_PER_DOC:
            target = next_available_filename(target)

    doc, table, target = ensure_document_and_table(target)
    ts = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    row = table.add_row()
    row.cells[0].text = f"{ts} — {context_text.strip() if context_text else '(no context provided)'}"
    p = row.cells[1].paragraphs[0]
    run = p.add_run()
    width = compute_column_image_width_inches(doc)
    run.add_picture(str(screenshot_path), width=Inches(width))
    doc.save(str(target))
    return target

HTML = """
<!doctype html>
<html>
<head>
  <meta charset="utf-8">
  <title>{{ title }}</title>
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <style>
    body { font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, Helvetica, Arial, sans-serif;
           background: #0b0b0c; color: #e7e7ea; display:flex; min-height:100vh; align-items:center; justify-content:center; }
    .card { width: 760px; max-width: 95vw; background: #17171a; border-radius: 20px; box-shadow: 0 10px 30px rgba(0,0,0,0.4);
            padding: 24px; }
    h1 { margin: 0 0 6px 0; font-size: 28px; }
    p.sub { margin: 0 0 18px 0; opacity: 0.8; }
    textarea, input[type=number] { width: 100%; background: #0f0f12; color: #e7e7ea; border: 1px solid #2a2a2e;
               border-radius: 12px; padding: 12px; font-size: 16px; outline: none; }
    textarea { height: 120px; }
    .grid { display:grid; grid-template-columns: 1fr 160px; gap: 12px; align-items:center; }
    .actions { display:flex; gap:12px; margin-top: 16px; align-items:center; flex-wrap:wrap; }
    .btn { display:inline-block; padding: 16px 24px; font-size: 22px; font-weight: 800; border:none; cursor:pointer;
           border-radius: 16px; transition: transform 0.05s ease; }
    .btn:active { transform: translateY(2px); }
    .btn-red { background:#c1121f; color:white; }
    .hint { opacity:0.7; font-size: 13px; }
    .flash { margin-top: 12px; padding: 10px 12px; border-radius: 10px; background: #12351f; color: #c7f7d1; }
    label { font-size: 14px; opacity: 0.9; }
  </style>
  <script>
    function onCaptureClick() {
      const btn = document.getElementById('captureBtn');
      const delay = document.getElementById('delay').value || '0';
      btn.disabled = true;
      btn.innerText = "Capture in " + delay + "s…";
      document.getElementById('form').submit();
    }
  </script>
</head>
<body>
  <div class="card">
    <h1>{{ title }}</h1>
    <p class="sub">Cross-platform (Windows/macOS/Linux). Click the big red button to capture and log to Word.</p>
    {% with messages = get_flashed_messages() %}
      {% if messages %}
        {% for m in messages %}
          <div class="flash">{{ m }}</div>
        {% endfor %}
      {% endif %}
    {% endwith %}
    <form id="form" method="POST" action="{{ url_for('capture') }}">
      <div class="grid">
        <div>
          <label for="context">Context</label>
          <textarea id="context" name="context" placeholder="What is the context?"></textarea>
        </div>
        <div>
          <label for="delay">Delay before capture (seconds)</label>
          <input id="delay" name="delay" type="number" step="0.1" min="0" max="60" value="{{ delay }}">
        </div>
      </div>
      <div class="actions">
        <button id="captureBtn" type="button" class="btn btn-red" onclick="onCaptureClick()">CAPTURE & LOG</button>
        <span class="hint">Use the delay to alt-tab to the target window before the shot.</span>
      </div>
    </form>
  </div>
</body>
</html>
"""

@app.get("/")
def index():
    try:
        delay = float(request.args.get("delay", DEFAULT_DELAY_SECONDS))
    except Exception:
        delay = DEFAULT_DELAY_SECONDS
    return render_template_string(HTML, title=APP_TITLE, delay=delay)

@app.post("/capture")
def capture():
    try:
        delay = float(request.form.get("delay", DEFAULT_DELAY_SECONDS))
    except Exception:
        delay = DEFAULT_DELAY_SECONDS
    delay = max(0.0, min(60.0, delay))

    context = request.form.get("context", "").strip()
    out_dir = user_documents_dir()
    out_dir.mkdir(parents=True, exist_ok=True)
    doc_path = out_dir / DEFAULT_DOC_NAME

    if delay > 0:
        time.sleep(delay)

    tmp_png = out_dir / f"context_shot_{int(time.time()*1000)}.png"
    try:
        take_full_screenshot_to(tmp_png)
        saved_to = append_entry(doc_path, context, tmp_png)
        flash(f"Saved to {saved_to} (delay={delay:.1f}s)")
    except Exception as e:
        flash(f"Error: {e}")
    finally:
        try:
            if tmp_png.exists():
                tmp_png.unlink()
        except Exception:
            pass

    return redirect(url_for('index', delay=f"{delay:.1f}"))

def main():
    url = f"http://{HOST}:{PORT}"
    webbrowser.open(url, new=2)
    app.run(host=HOST, port=PORT, debug=False)

if __name__ == "__main__":
    main()
