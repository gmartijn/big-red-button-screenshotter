#!/usr/bin/env python3
"""
Big Red Button — Cross-Platform Web Edition (Windows-friendly)
--------------------------------------------------------------
Now with optional Website Poller:
  - Enter a URL and an interval (seconds). The app will periodically take a headless
    screenshot of that page and append it to the Word log, with timestamp and URL.

Manual capture still works (big red button).

Tech:
  - Windows/Linux desktop screenshots: MSS
  - macOS desktop screenshots: `screencapture`
  - Website screenshots: Selenium (headless Chrome via webdriver-manager)
  - Word output: python-docx

Output file: ~/Documents/ContextShots.docx (auto-rotates after 90 rows)
"""
import os
import time
import shlex
import platform
import subprocess
import threading
from datetime import datetime
from pathlib import Path
import webbrowser

from flask import Flask, request, render_template_string, redirect, url_for, flash

# Word
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Screenshot (Windows/Linux desktop)
from mss import mss
from mss.tools import to_png

# Website screenshots (headless Chrome)
from selenium import webdriver
from selenium.webdriver.chrome.options import Options as ChromeOptions
from selenium.webdriver.chrome.service import Service as ChromeService
from webdriver_manager.chrome import ChromeDriverManager

APP_TITLE = "Big Red Button - Context + Screenshot (Cross-Platform)"
HOST = "127.0.0.1"
PORT = 8788
MAX_ROWS_PER_DOC = 90
DEFAULT_DOC_NAME = "ContextShots.docx"
DEFAULT_DELAY_SECONDS = 2.0

app = Flask(__name__)
app.secret_key = "context-shot-cross"

# Synchronize Word writes across threads
_doc_lock = threading.Lock()

# Poller state
_poller_thread = None
_poller_stop = threading.Event()
_poller_status = {"running": False, "url": "", "interval": 0.0}


# ----------------- Helpers -----------------

def user_documents_dir() -> Path:
    home = Path.home()
    docs = home / "Documents"
    return docs if docs.exists() else home

def next_available_filename(base: Path) -> Path:
    if not base.exists():
        return base
    stem, suffix = base.stem, base.suffix
    n = 2
    while True:
        cand = base.with_name(f"{stem} ({n}){suffix}")
        if not cand.exists():
            return cand
        n += 1

def _set_table_column_widths(table, widths_in_inches):
    for idx, width in enumerate(widths_in_inches):
        for cell in table.columns[idx].cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcW = OxmlElement('w:tcW')
            tcW.set(qn('w:type'), 'dxa')
            tcW.set(qn('w:w'), str(int(width * 1440)))  # 1 in = 1440 twips
            tcPr.append(tcW)

def _add_title_and_table(doc: Document, title_suffix: str = ""):
    title = doc.add_paragraph()
    run = title.add_run(f"Context + Screenshot Log {title_suffix}".strip())
    run.font.size = Pt(16)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    table = doc.add_table(rows=1, cols=2)
    hdr = table.rows[0].cells
    hdr[0].text = "Context (with timestamp)"
    hdr[1].text = "Screenshot"
    _set_table_column_widths(table, [3.1, 3.1])

def ensure_document_and_table(doc_path: Path):
    if doc_path.exists():
        doc = Document(str(doc_path))
        if doc.tables:
            table = doc.tables[0]
            if len(table.columns) != 2:
                doc = Document()
                _add_title_and_table(doc)
                return doc, doc.tables[0], doc_path
            rows = max(0, len(table.rows) - 1)
            if rows >= MAX_ROWS_PER_DOC:
                new_path = next_available_filename(doc_path)
                doc = Document()
                _add_title_and_table(doc, title_suffix=f"(Part {new_path.stem.split()[-1].strip('()')})")
                return doc, doc.tables[0], new_path
            return doc, table, doc_path
        else:
            _add_title_and_table(doc)
            return doc, doc.tables[0], doc_path
    else:
        doc = Document()
        _add_title_and_table(doc)
        return doc, doc.tables[0], doc_path

def compute_column_image_width_inches(doc: Document) -> float:
    section = doc.sections[0]
    page_width = section.page_width / 914400
    left = section.left_margin / 914400
    right = section.right_margin / 914400
    usable = page_width - left - right
    return max(2.2, min((usable / 2.0) - 0.15, 3.5))

def take_full_screenshot_to(file_path: Path) -> None:
    """Desktop full virtual screen capture (macOS via screencapture; others via MSS)."""
    system = platform.system().lower()
    if system == "darwin":
        cmd = f'screencapture -x -t png {shlex.quote(str(file_path))}'
        subprocess.run(cmd, shell=True, check=True)
        return
    with mss() as sct:
        monitor = sct.monitors[0]  # virtual screen (all monitors)
        raw = sct.grab(monitor)
        to_png(raw.rgb, raw.size, output=str(file_path))

def append_entry(doc_path: Path, context_text: str, screenshot_path: Path) -> Path:
    target = doc_path
    if target.exists():
        doc_tmp = Document(str(target))
        table_tmp = doc_tmp.tables[0] if doc_tmp.tables else None
        rows_tmp = (len(table_tmp.rows) - 1) if table_tmp else 0
        if rows_tmp >= MAX_ROWS_PER_DOC:
            target = next_available_filename(target)

    # Lock while writing the Word doc to avoid races with the poller
    with _doc_lock:
        doc, table, target = ensure_document_and_table(target)
        ts = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        row = table.add_row()
        row.cells[0].text = f"{ts} — {context_text.strip() if context_text else '(no context provided)'}"
        p = row.cells[1].paragraphs[0]
        run = p.add_run()
        width = compute_column_image_width_inches(doc)
        run.add_picture(str(screenshot_path), width=Inches(width))
        doc.save(str(target))
    return target

# ----------------- Website Screenshot (Selenium) -----------------

def _create_headless_driver():
    """Create a headless Chrome driver (robust across Chrome versions)."""
    last_err = None
    for headless_flag in ("--headless=new", "--headless"):
        try:
            options = ChromeOptions()
            options.add_argument(headless_flag)
            options.add_argument("--disable-gpu")
            options.add_argument("--no-sandbox")
            options.add_argument("--window-size=1920,1080")
            options.add_argument("--hide-scrollbars")
            options.add_argument("--disable-extensions")
            options.add_argument("--disable-dev-shm-usage")
            options.add_argument("--blink-settings=imagesEnabled=true")
            options.add_argument('--user-agent=Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/119 Safari/537.36')
            service = ChromeService(ChromeDriverManager().install())
            driver = webdriver.Chrome(service=service, options=options)
            driver.set_page_load_timeout(60)
            return driver
        except Exception as e:
            last_err = e
            continue
    raise RuntimeError(f"Failed to create headless Chrome driver: {last_err}")

def screenshot_website(url: str, out_path: Path, driver=None, full_page=True):
    """Take a screenshot of a website using a headless browser."""
    local_driver = driver or _create_headless_driver()
    own_driver = driver is None
    try:
        local_driver.get(url)
        # Wait a bit for dynamic content; try a trivial JS readyState wait
        try:
            for _ in range(50):
                rs = local_driver.execute_script("return document.readyState")
                if rs == "complete":
                    break
                time.sleep(0.1)
        except Exception:
            time.sleep(1.0)
        time.sleep(0.5)  # small settle

        if full_page:
            # Try to size to full page height (cap to a sane max)
            try:
                height = local_driver.execute_script(
                    "return Math.max(document.body.scrollHeight, document.documentElement.scrollHeight, "
                    "document.body.offsetHeight, document.documentElement.offsetHeight, "
                    "document.body.clientHeight, document.documentElement.clientHeight);"
                )
                height = int(max(1080, min(height or 1080, 20000)))
                local_driver.set_window_size(1920, height)
                time.sleep(0.2)
            except Exception:
                pass

        # Ensure PNG path
        out_path = out_path.with_suffix(".png")
        local_driver.save_screenshot(str(out_path))
    finally:
        if own_driver:
            try:
                local_driver.quit()
            except Exception:
                pass

# ----------------- Poller Thread -----------------

def _poller_loop(url: str, interval: float, doc_path: Path):
    """Background loop: screenshot URL every `interval` seconds and append to doc."""
    global _poller_status
    try:
        driver = _create_headless_driver()
    except Exception as e:
        _poller_status = {"running": False, "url": "", "interval": 0.0, "last_error": str(e), "last_saved": ""}
        return

    try:
        while not _poller_stop.is_set():
            tmp_png = doc_path.parent / f"auto_{int(time.time()*1000)}.png"
            last_saved = ""
            last_error = ""
            try:
                screenshot_website(url, tmp_png, driver=driver, full_page=True)
                ctx = f"Auto capture of {url}"
                last_saved = str(append_entry(doc_path, ctx, tmp_png))
            except Exception as e:
                last_error = str(e)
            finally:
                try:
                    if tmp_png.exists():
                        tmp_png.unlink()
                except Exception:
                    pass

            _poller_status = {
                "running": not _poller_stop.is_set(),
                "url": url,
                "interval": interval,
                "last_error": last_error,
                "last_saved": last_saved,
                "last_capture": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            }

            if _poller_stop.is_set():
                break
            if interval <= 0:
                break

            end_time = time.time() + interval
            while time.time() < end_time and not _poller_stop.is_set():
                time.sleep(0.25)
    finally:
        try:
            driver.quit()
        except Exception:
            pass
        _poller_status = {"running": False, "url": "", "interval": 0.0, "last_error": "", "last_saved": ""}

# ----------------- UI -----------------


HTML = """
<!doctype html>
<html>
<head>
  <meta charset="utf-8">
  <title>{{ title }}</title>
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <style>
    body { font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, Helvetica, Arial, sans-serif;
           background: #0b0b0c; color: #e7e7ea; display:flex; min-height:100vh; align-items:center; justify-content:center; }
    .wrap { width: 900px; max-width: 96vw; display:flex; flex-direction:column; gap:20px; }
    .card { background: #17171a; border-radius: 20px; box-shadow: 0 10px 30px rgba(0,0,0,0.4);
            padding: 24px; }
    h1 { margin: 0 0 6px 0; font-size: 28px; }
    p.sub { margin: 0 0 18px 0; opacity: 0.8; }
    textarea, input[type=number], input[type=text], input[type=url] {
      box-sizing: border-box;
      width: 100%; background: #0f0f12; color: #e7e7ea; border: 1px solid #2a2a2e;
      border-radius: 12px; padding: 12px; font-size: 16px; outline: none;
    }
    textarea { height: 120px; }
    .grid { display:grid; grid-template-columns: minmax(0, 1fr) 220px; gap: 12px; align-items:center; }
    .actions { display:flex; gap:12px; margin-top: 16px; align-items:center; flex-wrap:wrap; }
    .btn { display:inline-block; padding: 16px 24px; font-size: 22px; font-weight: 800; border:none; cursor:pointer;
           border-radius: 16px; transition: transform 0.05s ease; }
    .btn:active { transform: translateY(2px); }
    .btn-red { background:#c1121f; color:white; }
    .btn-green { background:#2d6a4f; color:white; }
    .btn-gray { background:#343a40; color:white; }
    .hint { opacity:0.7; font-size: 13px; }
    .flash { margin-top: 12px; padding: 10px 12px; border-radius: 10px; background: #12351f; color: #c7f7d1; }
    label { font-size: 14px; opacity: 0.9; }
    .status { margin-top: 8px; font-size: 14px; opacity: 0.85; }
  
    @media (max-width: 880px) {
      .grid { grid-template-columns: 1fr; }
    }
  </style>
  <script>
    function onManualCapture() {
      const btn = document.getElementById('captureBtn');
      const delay = document.getElementById('delay').value || '0';
      btn.disabled = true;
      btn.innerText = "Capture in " + delay + "s…";
      document.getElementById('manualForm').submit();
    }
  </script>
</head>
<body>
  <div class="wrap">
    <div class="card">
      <h1>{{ title }}</h1>
      <p class="sub">Manual desktop capture or website poller that auto-screenshots a URL on a fixed interval.</p>
      <p class="sub">Output document: <code>{{ doc_path }}</code></p>
      {% with messages = get_flashed_messages() %}
        {% if messages %}
          {% for m in messages %}
            <div class="flash">{{ m }}</div>
          {% endfor %}
        {% endif %}
      {% endwith %}
    </div>

    <div class="card">
      <h2>Manual Desktop Capture</h2>
      <form id="manualForm" method="POST" action="{{ url_for('capture') }}">
        <div class="grid">
          <div>
            <label for="context">Context</label>
            <textarea id="context" name="context" placeholder="What is the context?"></textarea>
          </div>
          <div>
            <label for="delay">Delay before capture (seconds)</label>
            <input id="delay" name="delay" type="number" step="0.1" min="0" max="60" value="{{ delay }}">
          </div>
        </div>
        <div class="actions">
          <button id="captureBtn" type="button" class="btn btn-red" onclick="onManualCapture()">CAPTURE & LOG</button>
          <span class="hint">Use the delay to alt-tab to the target window before the shot.</span>
        </div>
      </form>
    </div>

    <div class="card">
      <h2>Website Poller (Headless)</h2>
      <form method="POST" action="{{ url_for('poll_start') }}">
        <div class="grid">
          <div>
            <label for="poll_url">Website URL</label>
            <input id="poll_url" name="poll_url" type="url" placeholder="https://example.com" required value="{{ poll_url or '' }}">
          </div>
          <div>
            <label for="poll_interval">Interval seconds</label>
            <input id="poll_interval" name="poll_interval" type="number" step="1" min="5" max="86400" value="{{ poll_interval or 60 }}">
          </div>
        </div>
        <div class="actions">
          <button class="btn btn-green" type="submit">START POLLER</button>
          <a class="btn btn-gray" href="{{ url_for('poll_stop') }}">STOP POLLER</a>
          <span class="status">
            {% if poll_status.running %}
              ✅ Running every {{ poll_status.interval }}s on {{ poll_status.url }}<br/>
              🔄 Last capture: {{ poll_status.last_capture or "n/a" }}<br/>
              💾 Last saved: {{ poll_status.last_saved or "n/a" }}<br/>
              ⚠️ Last error: {{ poll_status.last_error or "none" }}
            {% else %}
              ⏸️ Poller is stopped
            {% endif %}
          </span>
        </div>
      </form>
      <p class="hint">The poller uses a headless browser and appends: "Auto capture of &lt;URL&gt;" with timestamp.</p>
        </div>

    <div class="card">
      <h2>Quick Test: One-off Website Screenshot</h2>
      <form method="POST" action="{{ url_for('poll_test') }}">
        <div class="grid">
          <div>
            <label for="test_url">Website URL</label>
            <input id="test_url" name="test_url" type="url" placeholder="https://example.com" required>
          </div>
          <div>
            <label>&nbsp;</label>
            <button class="btn btn-green" type="submit">TEST CAPTURE NOW</button>
          </div>
        </div>
        <p class="hint">Takes a single headless screenshot and appends it, without starting the poller.</p>
      </form>
    </div>
  </div>
</body>
</html>
"""

@app.get("/")
def index():
    try:
        delay = float(request.args.get("delay", DEFAULT_DELAY_SECONDS))
    except Exception:
        delay = DEFAULT_DELAY_SECONDS
    return render_template_string(HTML, title=APP_TITLE, doc_path=str(user_documents_dir() / DEFAULT_DOC_NAME),
        delay=delay,
        poll_status=_poller_status,
        poll_url=_poller_status.get("url", ""),
        poll_interval=int(_poller_status.get("interval", 60) or 60),
    )

@app.post("/capture")
def capture():
    try:
        delay = float(request.form.get("delay", DEFAULT_DELAY_SECONDS))
    except Exception:
        delay = DEFAULT_DELAY_SECONDS
    delay = max(0.0, min(60.0, delay))

    context = request.form.get("context", "").strip()
    out_dir = user_documents_dir()
    out_dir.mkdir(parents=True, exist_ok=True)
    doc_path = out_dir / DEFAULT_DOC_NAME

    if delay > 0:
        time.sleep(delay)

    tmp_png = out_dir / f"context_shot_{int(time.time()*1000)}.png"
    try:
        take_full_screenshot_to(tmp_png)
        saved_to = append_entry(doc_path, context, tmp_png)
        flash(f"Saved to {saved_to} (delay={delay:.1f}s)")
    except Exception as e:
        flash(f"Error: {e}")
    finally:
        try:
            if tmp_png.exists():
                tmp_png.unlink()
        except Exception:
            pass

    return redirect(url_for('index', delay=f"{delay:.1f}"))

@app.post("/poll/start")
def poll_start():
    global _poller_thread, _poller_stop, _poller_status
    url = (request.form.get("poll_url") or "").strip()
    try:
        interval = float(request.form.get("poll_interval", "60"))
    except Exception:
        interval = 60.0
    interval = max(5.0, min(86400.0, interval))

    if not url.lower().startswith(("http://", "https://")):
        flash("Please provide a valid URL starting with http:// or https://")
        return redirect(url_for('index'))

    if _poller_status["running"]:
        flash("Poller already running. Stop it first if you want to change URL/interval.")
        return redirect(url_for('index'))

    out_dir = user_documents_dir()
    out_dir.mkdir(parents=True, exist_ok=True)
    doc_path = out_dir / DEFAULT_DOC_NAME

    _poller_stop.clear()
    _poller_status = {"running": True, "url": url, "interval": interval}
    _poller_thread = threading.Thread(target=_poller_loop, args=(url, interval, doc_path), daemon=True)
    _poller_thread.start()
    flash(f"Poller started: every {interval:.0f}s on {url}")
    return redirect(url_for('index'))

@app.get("/poll/stop")
def poll_stop():
    global _poller_thread, _poller_stop, _poller_status
    if _poller_status["running"]:
        _poller_stop.set()
        th = _poller_thread
        _poller_thread = None
        _poller_status = {"running": False, "url": "", "interval": 0.0}
        if th is not None:
            th.join(timeout=5.0)
        flash("Poller stopped.")
    else:
        flash("Poller is not running.")
    return redirect(url_for('index'))

# Back-compat short routes
@app.post("/poll_start")
def poll_start_bc():
    return poll_start()

@app.get("/poll_stop")
def poll_stop_bc():
    return poll_stop()


@app.post("/poll/test")
def poll_test():
    url = (request.form.get("test_url") or "").strip()
    if not url.lower().startswith(("http://", "https://")):
        flash("Please provide a valid URL starting with http:// or https://")
        return redirect(url_for('index'))
    out_dir = user_documents_dir()
    out_dir.mkdir(parents=True, exist_ok=True)
    doc_path = out_dir / DEFAULT_DOC_NAME
    tmp_png = out_dir / f"test_{int(time.time()*1000)}.png"
    try:
        screenshot_website(url, tmp_png, driver=None, full_page=True)
        saved_to = append_entry(doc_path, f"Manual website capture of {url}", tmp_png)
        flash(f"Captured and saved to {saved_to}")
    except Exception as e:
        flash(f"Error during test capture: {e}")
    finally:
        try:
            if tmp_png.exists():
                tmp_png.unlink()
        except Exception:
            pass
    return redirect(url_for('index'))

def main():
    url = f"http://{HOST}:{PORT}"
    webbrowser.open(url, new=2)
    app.run(host=HOST, port=PORT, debug=False)

if __name__ == "__main__":
    main()